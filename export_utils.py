"""
Export utilities for generating DOCX and MD files.
"""

import io
import logging
from typing import Dict, Any
from datetime import datetime
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    import markdown
except ImportError:
    # markdown is optional for basic functionality
    markdown = None

logger = logging.getLogger(__name__)

def generate_docx(analysis_data: Dict[str, Any]) -> bytes:
    """
    Generate a DOCX file from analysis data.
    
    Args:
        analysis_data: Analysis data dictionary
        
    Returns:
        DOCX file as bytes
    """
    try:
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading(f"UN GA Analysis: {analysis_data['country']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run("Analysis Metadata:\n").bold = True
        metadata.add_run(f"Country/Entity: {analysis_data['country']}\n")
        metadata.add_run(f"Classification: {analysis_data['classification']}\n")
        metadata.add_run(f"Speech Date: {analysis_data.get('speech_date', 'Not specified')}\n")
        metadata.add_run(f"Analysis Date: {analysis_data['created_at'].strftime('%Y-%m-%d %H:%M:%S')}\n")
        if analysis_data.get('sdgs'):
            metadata.add_run(f"SDGs Mentioned: {', '.join([f'SDG {sdg}' for sdg in analysis_data['sdgs']])}\n")
        if analysis_data['classification'] == 'Development Partner':
            metadata.add_run(f"Africa Mentioned: {'Yes' if analysis_data['africa_mentioned'] else 'No'}\n")
        
        # Add separator
        doc.add_paragraph("=" * 80)
        
        # Add the analysis content
        if analysis_data.get('output_markdown'):
            # Convert markdown to formatted text
            _add_markdown_content(doc, analysis_data['output_markdown'])
        
        # Add raw text section
        if analysis_data.get('raw_text'):
            doc.add_heading("Original Speech Text", level=1)
            raw_text_para = doc.add_paragraph(analysis_data['raw_text'])
            raw_text_para.style = 'Normal'
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
        
    except Exception as e:
        logger.error(f"Failed to generate DOCX: {e}")
        raise Exception(f"Failed to generate DOCX file: {e}")

def generate_markdown(analysis_data: Dict[str, Any]) -> str:
    """
    Generate a markdown file from analysis data.
    
    Args:
        analysis_data: Analysis data dictionary
        
    Returns:
        Markdown content as string
    """
    try:
        md_content = []
        
        # Add title
        md_content.append(f"# UN GA Analysis: {analysis_data['country']}\n")
        
        # Add metadata
        md_content.append("## Analysis Metadata\n")
        md_content.append(f"- **Country/Entity:** {analysis_data['country']}")
        md_content.append(f"- **Classification:** {analysis_data['classification']}")
        md_content.append(f"- **Speech Date:** {analysis_data.get('speech_date', 'Not specified')}")
        md_content.append(f"- **Analysis Date:** {analysis_data['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
        
        if analysis_data.get('sdgs'):
            sdg_list = ', '.join([f'SDG {sdg}' for sdg in analysis_data['sdgs']])
            md_content.append(f"- **SDGs Mentioned:** {sdg_list}")
        
        if analysis_data['classification'] == 'Development Partner':
            africa_status = 'Yes' if analysis_data['africa_mentioned'] else 'No'
            md_content.append(f"- **Africa Mentioned:** {africa_status}")
        
        md_content.append("")  # Empty line
        
        # Add separator
        md_content.append("---\n")
        
        # Add the analysis content
        if analysis_data.get('output_markdown'):
            md_content.append(analysis_data['output_markdown'])
        
        # Add raw text section
        if analysis_data.get('raw_text'):
            md_content.append("\n## Original Speech Text\n")
            md_content.append(f"```\n{analysis_data['raw_text']}\n```")
        
        return '\n'.join(md_content)
        
    except Exception as e:
        logger.error(f"Failed to generate markdown: {e}")
        raise Exception(f"Failed to generate markdown file: {e}")

def _add_markdown_content(doc: docx.Document, markdown_text: str):
    """
    Add markdown content to a Word document with proper formatting.
    
    Args:
        doc: Word document object
        markdown_text: Markdown text to add
    """
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            
            # Handle bold text in bullets
            if '**' in bullet_text:
                _format_bold_text(para, bullet_text)
        
        # Handle numbered lists
        elif line and line[0].isdigit() and '. ' in line:
            bullet_text = line.split('. ', 1)[1]
            para = doc.add_paragraph(bullet_text, style='List Number')
            
            # Handle bold text in numbered items
            if '**' in bullet_text:
                _format_bold_text(para, bullet_text)
        
        # Regular paragraphs
        else:
            para = doc.add_paragraph(line)
            
            # Handle bold text in paragraphs
            if '**' in line:
                _format_bold_text(para, line)

def _format_bold_text(paragraph, text: str):
    """
    Format bold text in a paragraph.
    
    Args:
        paragraph: Word paragraph object
        text: Text containing **bold** markers
    """
    # Clear the paragraph first
    paragraph.clear()
    
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:  # Regular text
            paragraph.add_run(part)
        else:  # Bold text
            paragraph.add_run(part).bold = True

def format_filename(analysis_data: Dict[str, Any], extension: str) -> str:
    """
    Generate a clean filename for the analysis.
    
    Args:
        analysis_data: Analysis data dictionary
        extension: File extension (docx, md)
        
    Returns:
        Clean filename
    """
    country = analysis_data['country'].replace(' ', '_').replace('/', '_')
    date_str = analysis_data['created_at'].strftime('%Y%m%d_%H%M')
    classification = analysis_data['classification'].replace(' ', '_')
    
    return f"UN_GA_Analysis_{country}_{classification}_{date_str}.{extension}"

def create_export_files(analysis_data: Dict[str, Any]) -> Dict[str, bytes]:
    """
    Create both DOCX and Markdown export files.
    
    Args:
        analysis_data: Analysis data dictionary
        
    Returns:
        Dictionary with 'docx' and 'md' keys containing file bytes/content
    """
    try:
        exports = {}
        
        # Generate DOCX
        exports['docx'] = generate_docx(analysis_data)
        
        # Generate Markdown
        exports['md'] = generate_markdown(analysis_data)
        
        return exports
        
    except Exception as e:
        logger.error(f"Failed to create export files: {e}")
        raise Exception(f"Failed to create export files: {e}")

def get_file_size_mb(file_bytes: bytes) -> float:
    """
    Get file size in MB.
    
    Args:
        file_bytes: File bytes
        
    Returns:
        File size in MB
    """
    return len(file_bytes) / (1024 * 1024)
